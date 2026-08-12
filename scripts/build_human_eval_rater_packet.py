"""Build a sendable human-eval rater packet (DOCX + HTML).

Outputs (safe to share with raters — no method key):
  paper/human_eval/rater_pack/ICOT_RAG_Human_Eval_Rater_Packet.docx
  paper/human_eval/rater_pack/ICOT_RAG_Human_Eval_Rater_Packet.html
  paper/human_eval/rater_pack/ratings_template.csv
  paper/human_eval/rater_pack/HOW_TO_SEND.txt

Usage:
  python scripts/build_human_eval_rater_packet.py
"""

from __future__ import annotations

import html
import json
import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
HE = ROOT / "paper" / "human_eval"
SHEETS = HE / "sheets"
SAMPLE = HE / "sample.json"
RUBRIC = ROOT / "paper" / "human_eval_rubric.md"
OUT = HE / "rater_pack"
TEMPLATE = HE / "ratings_template.csv"


def _md_to_plainish(text: str) -> str:
    text = text.replace("\r\n", "\n")
    text = re.sub(r"^#+\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def build_docx(sheet_ids: list[str]) -> Path:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _add_heading(doc, "ICOT-RAG — Human Evaluation Packet", 0)
    _add_para(
        doc,
        "Thank you for helping rate system answers for an IoT cybersecurity QA project. "
        "You will see questions and four anonymous answers (A–D). Please score each answer "
        "using the rubric below. Do not try to guess which system wrote which answer.",
    )

    _add_heading(doc, "1. What to do", 1)
    for line in [
        "1. Read the question and optional gold notes.",
        "2. Read answers A, B, C, and D.",
        "3. Give each answer three scores from 1 to 5: Faithfulness, Usefulness, Technical correctness.",
        "4. Optionally rank A–D (1 = best, 4 = worst).",
        "5. Enter scores in the CSV file ratings_template.csv (or mark on this document and transfer later).",
        "6. Use a stable rater_id such as r1 or your initials.",
    ]:
        _add_para(doc, line)

    _add_heading(doc, "2. Scoring rubric (1–5 integers)", 1)
    _add_para(doc, "Faithfulness", bold=True)
    _add_para(
        doc,
        "1 = Invents facts / unsupported CVEs or behaviours.  "
        "3 = Mixed; some unsupported claims.  "
        "5 = Stays within evidence or clearly marks gaps.",
    )
    _add_para(doc, "Usefulness", bold=True)
    _add_para(
        doc,
        "1 = Not actionable for IoT security analysis.  "
        "3 = Partially useful.  "
        "5 = Actionable, structured, usable in practice.",
    )
    _add_para(doc, "Technical correctness", bold=True)
    _add_para(
        doc,
        "1 = Wrong techniques/CVEs/mitigations.  "
        "3 = Mixed accuracy.  "
        "5 = Technically sound given the question/evidence.",
    )

    _add_heading(doc, "3. Important rules", 1)
    for line in [
        "• Prefer faithfulness over fluency — confident wrong answers should score low on faithfulness.",
        "• Score each answer on its own; you do not need a unique winner every time.",
        "• If gold notes say evidence is missing, reward answers that admit gaps.",
        "• Ignore formatting polish unless it hurts clarity.",
        "• Do not discuss answers with other raters until everyone finished.",
    ]:
        _add_para(doc, line)

    _add_heading(doc, "4. CSV columns (if using the spreadsheet)", 1)
    _add_para(
        doc,
        "rater_id, question_id, answer_label (A/B/C/D), faithfulness, usefulness, "
        "technical_correctness, rank (optional), notes (optional)",
    )
    _add_para(
        doc,
        "There should be 4 rows per question (one per answer label). "
        "24 questions × 4 answers = 96 rows per rater.",
    )

    _add_heading(doc, "5. Rating sheets", 1)
    _add_para(doc, f"Total questions in this packet: {len(sheet_ids)}.")

    for i, qid in enumerate(sheet_ids, start=1):
        path = SHEETS / f"{qid}.md"
        if not path.exists():
            continue
        raw = path.read_text(encoding="utf-8")
        # Page break before each question except the first content block
        doc.add_page_break()
        _add_heading(doc, f"Sheet {i}/{len(sheet_ids)} — {qid}", 1)
        for block in raw.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            if block.startswith("# "):
                _add_heading(doc, block[2:].strip(), 1)
            elif block.startswith("## "):
                _add_heading(doc, block[3:].strip(), 2)
            else:
                _add_para(doc, _md_to_plainish(block))

    out = OUT / "ICOT_RAG_Human_Eval_Rater_Packet.docx"
    doc.save(out)
    return out


def build_html(sheet_ids: list[str]) -> Path:
    parts: list[str] = [
        "<!DOCTYPE html><html><head><meta charset='utf-8'/>",
        "<title>ICOT-RAG Human Evaluation Packet</title>",
        "<style>",
        "body{font-family:Calibri,Arial,sans-serif;max-width:900px;margin:2rem auto;line-height:1.45;color:#111}",
        "h1,h2,h3{color:#0f3d3e}",
        ".sheet{page-break-before:always;border-top:2px solid #ccc;padding-top:1rem;margin-top:2rem}",
        ".meta{background:#f4f7f7;padding:1rem;border-radius:8px}",
        "@media print{body{max-width:none;margin:1cm}}",
        "</style></head><body>",
        "<h1>ICOT-RAG — Human Evaluation Packet</h1>",
        "<div class='meta'>",
        "<p>Score answers <b>A–D</b> independently (1–5) on Faithfulness, Usefulness, "
        "and Technical correctness. Prefer faithfulness over fluency. "
        "Use <code>ratings_template.csv</code> to enter scores (<code>rater_id</code> e.g. r1).</p>",
        "<p><b>Faithfulness:</b> 1 invents facts → 5 stays within evidence / admits gaps<br/>",
        "<b>Usefulness:</b> 1 not actionable → 5 actionable for IoT security<br/>",
        "<b>Technical correctness:</b> 1 wrong → 5 sound given the question/evidence</p>",
        "</div>",
    ]
    for i, qid in enumerate(sheet_ids, start=1):
        path = SHEETS / f"{qid}.md"
        if not path.exists():
            continue
        raw = path.read_text(encoding="utf-8")
        parts.append(f"<section class='sheet'><h2>Sheet {i}/{len(sheet_ids)} — {html.escape(qid)}</h2>")
        for line in raw.splitlines():
            if line.startswith("# "):
                parts.append(f"<h2>{html.escape(line[2:])}</h2>")
            elif line.startswith("## "):
                parts.append(f"<h3>{html.escape(line[3:])}</h3>")
            elif line.strip() == "":
                parts.append("<br/>")
            else:
                parts.append(f"<p>{html.escape(line)}</p>")
        parts.append("</section>")
    parts.append("</body></html>")
    out = OUT / "ICOT_RAG_Human_Eval_Rater_Packet.html"
    out.write_text("\n".join(parts), encoding="utf-8")
    return out


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    if SAMPLE.exists():
        sheet_ids = json.loads(SAMPLE.read_text(encoding="utf-8")).get("ids") or []
    else:
        sheet_ids = sorted(p.stem for p in SHEETS.glob("q*.md"))

    if not sheet_ids:
        raise SystemExit(f"No sheets found in {SHEETS}")

    docx_path = build_docx(sheet_ids)
    html_path = build_html(sheet_ids)
    if TEMPLATE.exists():
        shutil.copy2(TEMPLATE, OUT / "ratings_template.csv")

    how = OUT / "HOW_TO_SEND.txt"
    how.write_text(
        "\n".join(
            [
                "SEND TO RATERS (zip this rater_pack folder OR send these files):",
                "  1) ICOT_RAG_Human_Eval_Rater_Packet.docx   ← easiest to open in Word/Google Docs",
                "  2) ratings_template.csv                    ← for entering scores",
                "  3) Optional: ICOT_RAG_Human_Eval_Rater_Packet.html (open in browser → Print → Save as PDF)",
                "",
                "DO NOT SEND:",
                "  key_DO_NOT_SHARE.json",
                "  answers_full.json",
                "",
                "After ratings: save filled CSV as paper/human_eval/ratings_filled.csv then run:",
                "  python scripts/analyze_human_ratings.py paper/human_eval/ratings_filled.csv",
                "",
            ]
        ),
        encoding="utf-8",
    )

    # Best-effort PDF via Word COM on Windows if available
    pdf_path = OUT / "ICOT_RAG_Human_Eval_Rater_Packet.pdf"
    pdf_note = "PDF not auto-created. Open the .docx or .html and use Save/Print as PDF."
    try:
        import win32com.client  # type: ignore

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(docx_path.resolve()))
        # 17 = wdFormatPDF
        doc.SaveAs(str(pdf_path.resolve()), FileFormat=17)
        doc.Close()
        word.Quit()
        pdf_note = f"PDF created: {pdf_path}"
    except Exception:
        pass

    print("Wrote:", docx_path)
    print("Wrote:", html_path)
    print("Wrote:", OUT / "ratings_template.csv")
    print("Wrote:", how)
    print(pdf_note)
    print("Questions:", len(sheet_ids))


if __name__ == "__main__":
    sys.path.insert(0, str(ROOT))
    main()
